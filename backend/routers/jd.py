import io
import re
from typing import Optional
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from openai import AsyncOpenAI
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from auth import get_current_user
from database import get_db
from models import JDRecord, User
from config import get_settings

settings = get_settings()
openai_client = AsyncOpenAI(api_key=settings.openai_api_key)
router = APIRouter(prefix="/api/jd", tags=["jd"])


class JDGenerateRequest(BaseModel):
    job_title: str
    location: Optional[str] = "Headquarters"
    department: Optional[str] = ""
    employment_type: Optional[str] = "Full-Time"
    salary_min: Optional[str] = ""
    salary_max: Optional[str] = ""
    reports_to: Optional[str] = ""
    notes: Optional[str] = ""
    use_web_context: bool = False
    current_jd: Optional[str] = ""


class JDExportRequest(BaseModel):
    content: str
    job_title: str
    location: Optional[str] = ""
    department: Optional[str] = ""
    employment_type: Optional[str] = ""


class JDRenameRequest(BaseModel):
    title: str


@router.post("/generate")
async def generate_jd(
    body: JDGenerateRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    salary_str = ""
    if body.salary_min and body.salary_max:
        salary_str = f"Salary Range: ${body.salary_min} – ${body.salary_max}"

    prompt = f"""Create a professional, comprehensive job description for the following role. Output ONLY clean plain text — do NOT use any markdown symbols such as #, ##, ###, *, **, -, --, bullets, or any other special formatting characters. Do not use asterisks, hashes, dashes as bullet points, or any markdown syntax whatsoever.

Job Title: {body.job_title}
Location: {body.location}
Department: {body.department}
Employment Type: {body.employment_type}
{salary_str}
Reports To: {body.reports_to}
Additional Notes from HR: {body.notes}
"""

    if body.current_jd:
        prompt += f"\n\nHere is the current/existing Job Description to use as a baseline or reference. Please enhance it using the new details provided above:\n{body.current_jd}\n\n"

    prompt += """Structure the job description with these clearly labeled sections. Write each section header in ALL CAPS followed by a colon and a new line. Use numbered items (1. 2. 3.) for lists instead of dashes or bullets. Keep single line spacing between items — do NOT add blank lines between list items. Add ONE blank line between sections only.

ROLE SUMMARY
Write 2-3 sentences summarizing the role.

KEY RESPONSIBILITIES
List 6-8 responsibilities as numbered items on consecutive lines with no blank lines between them.

REQUIREMENTS
List 5-7 requirements as numbered items on consecutive lines with no blank lines between them.

PREFERRED QUALIFICATIONS
List 3-4 items as numbered items on consecutive lines with no blank lines between them.

WHAT WE OFFER
List 4-5 benefits as numbered items on consecutive lines with no blank lines between them.

Write in a professional, engaging tone. Be specific and concrete. Do NOT use placeholders like [Company Name]. Always use "Superior Paving Corp." as the company name. If any detail is missing, make up a reasonable professional default or omit it naturally, ensuring the document is completely finished with no bracketed placeholders left for the user to edit. Do NOT use any markdown formatting characters."""

    response = await openai_client.chat.completions.create(
        model=settings.openai_chat_model,
        messages=[
            {"role": "system", "content": "You are an expert HR professional and talent acquisition specialist who writes compelling, inclusive job descriptions."},
            {"role": "user", "content": prompt},
        ],
        temperature=0.7,
        max_tokens=1500,
    )

    content = response.choices[0].message.content
    content = re.sub(r'(?i)\[company name\]', 'Superior Paving Corp.', content)

    # Save record
    record = JDRecord(
        user_id=current_user.id,
        title=body.job_title,
        department=body.department,
        content=content,
        form_data=body.model_dump(),
    )
    db.add(record)
    await db.commit()
    await db.refresh(record)

    return {
        "id": record.id,
        "content": content,
        "job_title": body.job_title,
        "department": body.department,
        "location": body.location,
        "employment_type": body.employment_type,
    }


def _set_para_spacing(para, before_pt=0, after_pt=3, line_spacing_pt=14):
    """Set tight paragraph spacing to avoid large gaps."""
    pPr = para._p.get_or_add_pPr()
    pSpacing = OxmlElement('w:spacing')
    pSpacing.set(qn('w:before'), str(int(before_pt * 20)))
    pSpacing.set(qn('w:after'), str(int(after_pt * 20)))
    pSpacing.set(qn('w:line'), str(int(line_spacing_pt * 20)))
    pSpacing.set(qn('w:lineRule'), 'exact')
    pPr.append(pSpacing)


def _strip_markdown(text: str) -> str:
    """Remove any residual markdown symbols from text."""
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'^[-*]\s+', '', text, flags=re.MULTILINE)
    return text.strip()


def _add_plaintext_to_docx(doc: DocxDocument, content: str):
    """Parse plain-text content (ALL CAPS headers, numbered lists, paragraphs) into a DOCX."""
    cleaned = _strip_markdown(content)
    lines = cleaned.split("\n")

    for line in lines:
        stripped = line.strip()

        # Skip blank lines — we control spacing via paragraph settings
        if not stripped:
            continue

        # ALL CAPS section header (e.g. ROLE SUMMARY, KEY RESPONSIBILITIES)
        if re.match(r'^[A-Z][A-Z\s]{3,}$', stripped):
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
            run.font.name = 'Arial'
            _set_para_spacing(p, before_pt=10, after_pt=3, line_spacing_pt=14)
            # Add a bottom border (underline effect)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'C8E6C9')
            pBdr.append(bottom)
            pPr.append(pBdr)
            continue

        # Numbered list item (1. 2. 3.)
        numbered_match = re.match(r'^(\d+\.\s+)(.*)', stripped)
        if numbered_match:
            p = doc.add_paragraph()
            num_run = p.add_run(numbered_match.group(1))
            num_run.bold = True
            num_run.font.size = Pt(10)
            num_run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
            num_run.font.name = 'Arial'
            text_run = p.add_run(numbered_match.group(2))
            text_run.font.size = Pt(10.5)
            text_run.font.name = 'Calibri'
            _set_para_spacing(p, before_pt=0, after_pt=2, line_spacing_pt=14)
            p.paragraph_format.left_indent = Cm(0.3)
            continue

        # Regular paragraph / body text
        p = doc.add_paragraph(stripped)
        p.runs[0].font.size = Pt(10.5)
        p.runs[0].font.name = 'Calibri'
        _set_para_spacing(p, before_pt=0, after_pt=4, line_spacing_pt=14)


@router.post("/export-docx")
async def export_docx(body: JDExportRequest):
    doc = DocxDocument()

    # Add company logo
    try:
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_run = logo_para.add_run()
        logo_run.add_picture("logo.png", width=Inches(3.0))
    except Exception as e:
        print(f"Could not load logo: {e}")

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(body.job_title)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

    # Meta info
    meta_parts = []
    if body.department:
        meta_parts.append(body.department)
    if body.location:
        meta_parts.append(body.location)
    if body.employment_type:
        meta_parts.append(body.employment_type)
    if meta_parts:
        meta_para = doc.add_paragraph(" | ".join(meta_parts))
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    _add_plaintext_to_docx(doc, body.content)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    safe_name = body.job_title.replace(" ", "_").lower()
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{safe_name}_jd.docx"'},
    )


@router.get("")
async def list_jds(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(JDRecord)
        .where(JDRecord.user_id == current_user.id)
        .order_by(JDRecord.created_at.desc())
    )
    records = result.scalars().all()
    return [
        {
            "id": r.id,
            "title": r.title,
            "department": r.department,
            "content": r.content,
            "form_data": r.form_data,
            "created_at": r.created_at.isoformat() if r.created_at else None,
        }
        for r in records
    ]


@router.put("/{record_id}")
async def rename_jd(
    record_id: str,
    body: JDRenameRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(JDRecord).where(
            JDRecord.id == record_id,
            JDRecord.user_id == current_user.id,
        )
    )
    record = result.scalar_one_or_none()
    if not record:
        raise HTTPException(status_code=404, detail="JD record not found")
    record.title = body.title
    await db.commit()
    return {"ok": True}


@router.delete("/{record_id}")
async def delete_jd(
    record_id: str,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(JDRecord).where(
            JDRecord.id == record_id,
            JDRecord.user_id == current_user.id,
        )
    )
    record = result.scalar_one_or_none()
    if not record:
        raise HTTPException(status_code=404, detail="JD record not found")
    await db.delete(record)
    await db.commit()
    return {"ok": True}
