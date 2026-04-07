import io
import re
from typing import Optional
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from openai import AsyncOpenAI
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from auth import get_current_user
from database import get_db
from models import OfferLetterRecord, User
from config import get_settings

settings = get_settings()
openai_client = AsyncOpenAI(api_key=settings.openai_api_key)
router = APIRouter(prefix="/api/offer", tags=["offer"])


class OfferGenerateRequest(BaseModel):
    candidate_name: str
    position: str
    department: Optional[str] = ""
    start_date: Optional[str] = ""
    salary: Optional[str] = ""
    manager_name: Optional[str] = ""
    location: Optional[str] = "Headquarters"
    employment_type: Optional[str] = "Full-Time"
    benefits: Optional[str] = ""
    additional_notes: Optional[str] = ""


class OfferExportRequest(BaseModel):
    content: str
    candidate_name: str
    position: str


class OfferRenameRequest(BaseModel):
    candidate_name: Optional[str] = None
    position: Optional[str] = None


@router.post("/generate")
async def generate_offer_letter(
    body: OfferGenerateRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    prompt = f"""Write a professional, warm, and legally sound employment offer letter. Output ONLY clean plain text — do NOT use any markdown symbols such as #, ##, ###, *, **, -, or any other special formatting characters. Do not use asterisks, hashes, dashes as bullet points, or any markdown syntax whatsoever.

Candidate Name: {body.candidate_name}
Position: {body.position}
Department: {body.department}
Start Date: {body.start_date}
Compensation: {body.salary}
Reporting To: {body.manager_name}
Location: {body.location}
Employment Type: {body.employment_type}
Benefits: {body.benefits}
Additional Notes: {body.additional_notes}

Write the offer letter as a complete formal business letter with these sections in order. Do NOT add blank lines between consecutive paragraphs within a section. Add ONE blank line between major sections only. Do NOT use bullet points or numbered lists — write everything in full paragraph sentences.

Write the letter with:
1. A formal date line and salutation (Dear [Candidate Name],)
2. An opening congratulatory paragraph welcoming them to the team
3. A position details paragraph covering job title, department, start date, and location
4. A compensation and benefits paragraph describing salary and key benefits
5. An employment terms paragraph outlining the nature of employment
6. A next steps paragraph with signing deadline and required documents
7. A warm closing with signature block

Use placeholders like [Company Name] and [HR Manager Name] where appropriate. Write in a professional but welcoming tone. Do NOT use any markdown formatting characters."""

    response = await openai_client.chat.completions.create(
        model=settings.openai_chat_model,
        messages=[
            {"role": "system", "content": "You are an experienced HR professional writing formal employment offer letters."},
            {"role": "user", "content": prompt},
        ],
        temperature=0.5,
        max_tokens=1500,
    )

    content = response.choices[0].message.content

    record = OfferLetterRecord(
        user_id=current_user.id,
        candidate_name=body.candidate_name,
        position=body.position,
        content=content,
        form_data=body.model_dump(),
    )
    db.add(record)
    await db.commit()
    await db.refresh(record)

    return {"id": record.id, "content": content, "candidate_name": body.candidate_name, "position": body.position}


def _set_para_spacing(para, before_pt=0, after_pt=3, line_spacing_pt=14):
    """Set tight paragraph spacing to avoid large gaps."""
    pPr = para._p.get_or_add_pPr()
    pSpacing = OxmlElement('w:spacing')
    pSpacing.set(qn('w:before'), str(int(before_pt * 20)))
    pSpacing.set(qn('w:after'), str(int(after_pt * 20)))
    pSpacing.set(qn('w:line'), str(int(line_spacing_pt * 20)))
    pSpacing.set(qn('w:lineRule'), 'exact')
    pPr.append(pSpacing)


def _strip_markdown(text: str) -> str:
    """Remove any residual markdown symbols from text."""
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'^[-*]\s+', '', text, flags=re.MULTILINE)
    return text.strip()


def _add_plaintext_to_docx(doc: DocxDocument, content: str):
    """Parse plain-text content (paragraphs, numbered items, ALL CAPS headers) into a DOCX."""
    cleaned = _strip_markdown(content)
    lines = cleaned.split("\n")

    for line in lines:
        stripped = line.strip()

        # Skip blank lines — spacing handled via paragraph settings
        if not stripped:
            continue

        # ALL CAPS section header
        if re.match(r'^[A-Z][A-Z\s]{3,}$', stripped):
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
            run.font.name = 'Arial'
            _set_para_spacing(p, before_pt=10, after_pt=3, line_spacing_pt=14)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'C8E6C9')
            pBdr.append(bottom)
            pPr.append(pBdr)
            continue

        # Numbered list item (1. 2. 3.)
        numbered_match = re.match(r'^(\d+\.\s+)(.*)', stripped)
        if numbered_match:
            p = doc.add_paragraph()
            num_run = p.add_run(numbered_match.group(1))
            num_run.bold = True
            num_run.font.size = Pt(10)
            num_run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
            num_run.font.name = 'Arial'
            text_run = p.add_run(numbered_match.group(2))
            text_run.font.size = Pt(10.5)
            text_run.font.name = 'Calibri'
            _set_para_spacing(p, before_pt=0, after_pt=2, line_spacing_pt=14)
            p.paragraph_format.left_indent = Cm(0.3)
            continue

        # Regular paragraph / body text
        p = doc.add_paragraph(stripped)
        p.runs[0].font.size = Pt(10.5)
        p.runs[0].font.name = 'Calibri'
        _set_para_spacing(p, before_pt=0, after_pt=5, line_spacing_pt=15)


@router.post("/export-docx")
async def export_offer_docx(body: OfferExportRequest):
    doc = DocxDocument()

    # Add company logo
    try:
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_run = logo_para.add_run()
        logo_run.add_picture("logo.png", width=Inches(3.0))
    except Exception as e:
        print(f"Could not load logo: {e}")

    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run("[Company Name]")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

    doc.add_paragraph()
    title_para = doc.add_paragraph()
    r = title_para.add_run("OFFER OF EMPLOYMENT")
    r.bold = True
    r.font.size = Pt(14)

    doc.add_paragraph()

    _add_plaintext_to_docx(doc, body.content)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    safe_name = body.candidate_name.replace(" ", "_").lower()
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="offer_{safe_name}.docx"'},
    )


@router.get("")
async def list_offers(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(OfferLetterRecord)
        .where(OfferLetterRecord.user_id == current_user.id)
        .order_by(OfferLetterRecord.created_at.desc())
    )
    records = result.scalars().all()
    return [
        {
            "id": r.id,
            "candidate_name": r.candidate_name,
            "position": r.position,
            "content": r.content,
            "form_data": r.form_data,
            "created_at": r.created_at.isoformat(),
        }
        for r in records
    ]


@router.put("/{record_id}")
async def rename_offer(
    record_id: str,
    body: OfferRenameRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(OfferLetterRecord).where(
            OfferLetterRecord.id == record_id,
            OfferLetterRecord.user_id == current_user.id,
        )
    )
    record = result.scalar_one_or_none()
    if not record:
        raise HTTPException(status_code=404, detail="Offer letter not found")
    if body.candidate_name is not None:
        record.candidate_name = body.candidate_name
    if body.position is not None:
        record.position = body.position
    await db.commit()
    return {"ok": True}


@router.delete("/{record_id}")
async def delete_offer(
    record_id: str,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(OfferLetterRecord).where(
            OfferLetterRecord.id == record_id,
            OfferLetterRecord.user_id == current_user.id,
        )
    )
    record = result.scalar_one_or_none()
    if not record:
        raise HTTPException(status_code=404, detail="Offer letter not found")
    await db.delete(record)
    await db.commit()
    return {"ok": True}
